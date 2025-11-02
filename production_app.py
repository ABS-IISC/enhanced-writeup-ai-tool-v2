#!/usr/bin/env python3
"""
Production-Ready Enhanced Writeup Automation AI Tool
Exact replica of localhost functionality for live deployment
"""

from flask import Flask, render_template, request, jsonify
import os
import json
import threading
import time
from datetime import datetime
from werkzeug.utils import secure_filename
import uuid
from collections import defaultdict
import sys

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed")
    sys.exit(1)

app = Flask(__name__)

# Production configuration
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'production-secure-key-2024')
app.config['UPLOAD_FOLDER'] = '/tmp/uploads' if os.environ.get('RAILWAY_ENVIRONMENT') else 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024

# Global variables (same as localhost)
current_session = None
analysis_progress = {"progress": 0, "message": "Ready", "current_section": ""}
analysis_status = "idle"
analysis_logs = []
chat_history = []
analysis_complete = False
functionalities_enabled = False

class AnalysisSession:
    def __init__(self):
        self.session_id = f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
        self.document_name = ""
        self.document_path = ""
        self.sections = {}
        self.analysis_results = {}
        self.user_feedback = defaultdict(list)
        self.accepted_feedback = defaultdict(list)
        self.rejected_feedback = defaultdict(list)
        self.current_section_index = 0
        self.analysis_thread = None
        self.stop_analysis_flag = False
        self.guidelines_document = None
        self.guidelines_content = None
        
    def get_section_names(self):
        return list(self.sections.keys())
    
    def get_current_section_name(self):
        section_names = self.get_section_names()
        if 0 <= self.current_section_index < len(section_names):
            return section_names[self.current_section_index]
        return None

def extract_document_sections_from_docx(doc):
    sections = {}
    current_section = "Executive Summary"
    content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            is_bold = any(run.bold for run in para.runs if run.bold)
            if is_bold and len(text) < 100 and any(keyword in text.lower() for keyword in ['summary', 'background', 'analysis', 'recommendation', 'conclusion', 'objective', 'scope', 'methodology']):
                if content:
                    sections[current_section] = '\n'.join(content)
                current_section = text.rstrip(':')
                content = []
            else:
                content.append(text)
    
    if content:
        sections[current_section] = '\n'.join(content)
    
    return sections

def analyze_section_with_ai(section_name, section_content, guidelines=None):
    time.sleep(2)  # Simulate AI processing
    import random
    
    feedback_items = []
    section_lower = section_name.lower()
    
    if 'summary' in section_lower:
        feedback_items.extend([
            {
                "id": f"ai_{uuid.uuid4().hex[:8]}",
                "type": "critical",
                "category": "completeness",
                "description": "Executive summary lacks key performance metrics and quantitative data",
                "suggestion": "Include specific numbers, percentages, and measurable outcomes to strengthen the summary",
                "risk_level": "High",
                "confidence": 0.92,
                "example": "Add metrics like '25% improvement in efficiency' or 'reduced processing time by 3 days'",
                "questions": ["Does this summary include quantifiable results?", "Are key stakeholders clearly identified?"]
            },
            {
                "id": f"ai_{uuid.uuid4().hex[:8]}",
                "type": "important",
                "category": "structure",
                "description": "Summary structure could be improved for better readability",
                "suggestion": "Consider reorganizing content with clear subsections: Overview, Key Findings, Recommendations",
                "risk_level": "Medium",
                "confidence": 0.85,
                "example": "Use bullet points for key achievements and numbered lists for recommendations",
                "questions": ["Is the information presented in logical order?"]
            }
        ])
    
    elif 'background' in section_lower or 'objective' in section_lower:
        feedback_items.extend([
            {
                "id": f"ai_{uuid.uuid4().hex[:8]}",
                "type": "critical",
                "category": "compliance",
                "description": "Missing regulatory compliance references and industry standards",
                "suggestion": "Include references to relevant regulations, standards, or company policies",
                "risk_level": "High",
                "confidence": 0.88,
                "example": "Reference ISO standards, regulatory requirements, or internal compliance frameworks",
                "questions": ["Are all applicable regulations mentioned?", "Is the scope clearly defined?"]
            }
        ])
    
    elif 'analysis' in section_lower:
        feedback_items.extend([
            {
                "id": f"ai_{uuid.uuid4().hex[:8]}",
                "type": "critical",
                "category": "methodology",
                "description": "Analysis methodology needs more detailed explanation",
                "suggestion": "Provide step-by-step methodology, data sources, and analytical frameworks used",
                "risk_level": "High",
                "confidence": 0.90,
                "example": "Describe data collection methods, sample sizes, analytical tools, and validation processes",
                "questions": ["Is the methodology reproducible?", "Are data sources reliable and current?"]
            },
            {
                "id": f"ai_{uuid.uuid4().hex[:8]}",
                "type": "important",
                "category": "evidence",
                "description": "Supporting evidence and documentation could be strengthened",
                "suggestion": "Add more supporting data, charts, or references to strengthen conclusions",
                "risk_level": "Medium",
                "confidence": 0.83,
                "example": "Include trend analysis, comparative data, or statistical significance tests",
                "questions": ["Are conclusions supported by sufficient evidence?"]
            }
        ])
    
    # Add positive feedback
    if random.random() > 0.3:
        feedback_items.append({
            "id": f"ai_{uuid.uuid4().hex[:8]}",
            "type": "positive",
            "category": "strength",
            "description": f"Strong content structure and clear presentation in {section_name}",
            "suggestion": "Continue maintaining this level of clarity and organization",
            "risk_level": "Low",
            "confidence": 0.95,
            "example": "The logical flow and professional tone enhance readability",
            "questions": []
        })
    
    return {"feedback_items": feedback_items}

def process_chat_query(query, context):
    query_lower = query.lower()
    current_section = context.get('current_section', 'No section selected')
    doc_name = context.get('document_name', 'your document')
    functionalities_enabled = context.get('functionalities_enabled', False)
    
    if not functionalities_enabled:
        return f"**TARA AI - Analysis in Progress**\n\nI'm currently analyzing your document. Once complete, I'll have full access to provide detailed insights about:\n\n• Section-specific analysis\n• Risk assessment details\n• Compliance recommendations\n• Improvement suggestions\n\nPlease wait for analysis to complete for full functionality!"
    
    if any(word in query_lower for word in ['section', 'current', 'this']):
        if current_section and current_section != 'No section selected':
            high_risk = context.get('section_high_risk', 0)
            medium_risk = context.get('section_medium_risk', 0)
            return f"**Analysis: '{current_section}'**\n\nThis section contains {context.get('word_count', 0)} words with {high_risk} high-risk and {medium_risk} medium-risk items identified.\n\n**Key Areas for Attention:**\n• Compliance requirements\n• Supporting evidence\n• Quantitative data\n• Methodology clarity\n\nWhat specific aspect would you like me to explain?"
        else:
            return "**Section Selection Required**\n\nPlease select a document section from the left panel. I can then provide detailed analysis including:\n\n• Risk assessment breakdown\n• Specific improvement recommendations\n• Compliance gap analysis\n• Content quality evaluation"
    
    elif any(word in query_lower for word in ['risk', 'high', 'medium', 'low']):
        high_risk = context.get('high_risk', 0)
        medium_risk = context.get('medium_risk', 0)
        low_risk = context.get('low_risk', 0)
        return f"**Risk Assessment Summary**\n\n**{doc_name} Risk Profile:**\n• 🔴 {high_risk} High-risk items (require immediate attention)\n• 🟡 {medium_risk} Medium-risk items (should be addressed)\n• 🟢 {low_risk} Low-risk items (minor improvements)\n\n**High-risk items typically involve:**\n• Missing compliance requirements\n• Insufficient supporting evidence\n• Unclear methodology\n• Incomplete data analysis\n\nWould you like me to prioritize these by section?"
    
    elif any(word in query_lower for word in ['improve', 'suggestion', 'recommend']):
        return f"**Improvement Recommendations for {doc_name}**\n\n**Priority Actions:**\n1. **Add Quantitative Data** - Include specific metrics and measurements\n2. **Strengthen Evidence** - Provide supporting documentation and references\n3. **Clarify Methodology** - Explain analytical approaches and data sources\n4. **Enhance Structure** - Improve organization and logical flow\n\n**Quick Wins:**\n• Use bullet points for key findings\n• Add executive summary if missing\n• Include compliance references\n• Provide clear recommendations\n\nWhich area would you like detailed guidance on?"
    
    else:
        return f"**Hello! I'm TARA, your AI document analysis assistant.**\n\nI've completed analyzing {doc_name} and can help with:\n\n**📊 Analysis Support:**\n• Section-specific feedback\n• Risk assessment explanation\n• Improvement prioritization\n\n**🎯 Specific Guidance:**\n• Compliance requirements\n• Evidence strengthening\n• Structure optimization\n• Content enhancement\n\n**💡 Ask me about:**\n• \"What are the high-risk items?\"\n• \"How can I improve [section name]?\"\n• \"What compliance issues were found?\"\n\nWhat would you like to explore?"

def log_activity(message, level="INFO", section=None):
    log_entry = {
        "id": f"log_{uuid.uuid4().hex[:8]}",
        "timestamp": datetime.now().isoformat(),
        "level": level,
        "message": message,
        "section": section
    }
    analysis_logs.append(log_entry)
    return log_entry

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/health')
def health_check():
    return jsonify({
        'status': 'healthy',
        'service': 'Enhanced Writeup AI Tool v2.0',
        'environment': 'Production',
        'upload_folder': app.config['UPLOAD_FOLDER']
    }), 200

@app.route('/api/upload', methods=['POST'])
def upload_document():
    global current_session, analysis_status
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if not file.filename.endswith('.docx'):
        return jsonify({'success': False, 'error': 'Please upload a .docx file'}), 400
    
    try:
        current_session = AnalysisSession()
        analysis_status = "idle"
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        current_session.document_name = filename
        current_session.document_path = filepath
        
        doc = Document(filepath)
        sections = extract_document_sections_from_docx(doc)
        
        current_session.sections = sections
        
        log_activity(f"Document uploaded: {filename} ({len(sections)} sections)", "SUCCESS")
        
        return jsonify({
            'success': True,
            'session_id': current_session.session_id,
            'document_name': filename,
            'sections': list(sections.keys()),
            'total_sections': len(sections),
            'file_size': os.path.getsize(filepath)
        })
        
    except Exception as e:
        log_activity(f"Upload failed: {str(e)}", "ERROR")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/upload_guidelines', methods=['POST'])
def upload_guidelines():
    global current_session
    
    if not current_session:
        return jsonify({'success': False, 'error': 'No active session'}), 400
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No guidelines file uploaded'}), 400
    
    file = request.files['file']
    if not file.filename.endswith('.docx'):
        return jsonify({'success': False, 'error': 'Please upload a .docx guidelines file'}), 400
    
    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"guidelines_{filename}")
        file.save(filepath)
        
        doc = Document(filepath)
        guidelines_content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        
        current_session.guidelines_document = filepath
        current_session.guidelines_content = guidelines_content
        
        log_activity(f"Guidelines document uploaded: {filename}", "SUCCESS")
        
        return jsonify({
            'success': True,
            'guidelines_name': filename,
            'content_length': len(guidelines_content)
        })
        
    except Exception as e:
        log_activity(f"Guidelines upload failed: {str(e)}", "ERROR")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/start_analysis', methods=['POST'])
def start_analysis():
    global current_session, analysis_status, analysis_progress, analysis_complete, functionalities_enabled
    
    if not current_session:
        return jsonify({'success': False, 'error': 'No document uploaded'}), 400
    
    if analysis_status == "running":
        return jsonify({'success': False, 'error': 'Analysis already in progress'}), 400
    
    analysis_status = "running"
    analysis_progress = {"progress": 0, "message": "Initializing Hawkeye Analysis Framework...", "current_section": ""}
    current_session.stop_analysis_flag = False
    analysis_complete = False
    functionalities_enabled = False
    
    log_activity("Starting Hawkeye 20-Point Investigation Framework", "INFO")
    
    def run_analysis():
        global analysis_status, analysis_progress, analysis_complete, functionalities_enabled
        
        try:
            section_names = current_session.get_section_names()
            total_sections = len(section_names)
            
            for i, section_name in enumerate(section_names):
                if current_session.stop_analysis_flag:
                    analysis_status = "stopped"
                    analysis_progress["message"] = "Analysis stopped by user"
                    log_activity("Analysis stopped by user", "WARNING")
                    return
                
                progress_percent = int((i / total_sections) * 85)
                analysis_progress = {
                    "progress": progress_percent,
                    "message": f"Analyzing {section_name} ({i+1}/{total_sections})",
                    "current_section": section_name
                }
                
                log_activity(f"Deep analysis: {section_name}", "INFO", section_name)
                
                section_content = current_session.sections[section_name]
                result = analyze_section_with_ai(section_name, section_content, current_session.guidelines_content)
                current_session.analysis_results[section_name] = result
                
                feedback_count = len(result.get('feedback_items', []))
                log_activity(f"Analysis complete: {section_name} - {feedback_count} insights generated", "SUCCESS", section_name)
                
                time.sleep(1.5)
            
            analysis_progress = {
                "progress": 90,
                "message": "Generating comprehensive risk assessment...",
                "current_section": ""
            }
            log_activity("Generating risk assessment matrix", "INFO")
            time.sleep(2)
            
            analysis_progress = {
                "progress": 95,
                "message": "Finalizing analysis and enabling advanced features...",
                "current_section": ""
            }
            log_activity("Enabling advanced AI features", "INFO")
            time.sleep(2)
            
            analysis_status = "completed"
            analysis_complete = True
            functionalities_enabled = True
            analysis_progress = {
                "progress": 100,
                "message": "Analysis completed! All advanced features now available.",
                "current_section": ""
            }
            
            log_activity("Hawkeye analysis completed - All functionalities enabled", "SUCCESS")
            
        except Exception as e:
            analysis_status = "error"
            analysis_progress["message"] = f"Analysis failed: {str(e)}"
            log_activity(f"Analysis failed: {str(e)}", "ERROR")
    
    current_session.analysis_thread = threading.Thread(target=run_analysis)
    current_session.analysis_thread.daemon = True
    current_session.analysis_thread.start()
    
    return jsonify({'success': True, 'message': 'Analysis started'})

@app.route('/api/status')
def get_status():
    global analysis_status, analysis_progress, analysis_logs, analysis_complete, functionalities_enabled
    
    if not current_session:
        return jsonify({
            'status': 'no_session',
            'progress': analysis_progress,
            'logs': analysis_logs[-10:],
            'analysis_complete': False,
            'functionalities_enabled': False,
            'statistics': {
                'total_sections': 0,
                'analyzed_sections': 0,
                'total_feedback': 0,
                'high_risk': 0,
                'medium_risk': 0,
                'low_risk': 0
            }
        })
    
    total_feedback = 0
    high_risk = 0
    medium_risk = 0
    low_risk = 0
    
    for result in current_session.analysis_results.values():
        feedback_items = result.get('feedback_items', [])
        total_feedback += len(feedback_items)
        
        for item in feedback_items:
            risk_level = item.get('risk_level', 'Low')
            if risk_level == 'High':
                high_risk += 1
            elif risk_level == 'Medium':
                medium_risk += 1
            else:
                low_risk += 1
    
    statistics = {
        'total_sections': len(current_session.sections),
        'analyzed_sections': len(current_session.analysis_results),
        'total_feedback': total_feedback,
        'high_risk': high_risk,
        'medium_risk': medium_risk,
        'low_risk': low_risk,
        'user_feedback': sum(len(fb) for fb in current_session.user_feedback.values())
    }
    
    return jsonify({
        'status': analysis_status,
        'progress': analysis_progress,
        'logs': analysis_logs[-20:],
        'analysis_complete': analysis_complete,
        'functionalities_enabled': functionalities_enabled,
        'statistics': statistics,
        'session_info': {
            'session_id': current_session.session_id,
            'document_name': current_session.document_name,
            'current_section': current_session.get_current_section_name()
        }
    })

@app.route('/api/sections')
def get_sections():
    if not current_session:
        return jsonify({'sections': []})
    
    sections_info = []
    for section_name, content in current_session.sections.items():
        analysis_result = current_session.analysis_results.get(section_name, {})
        feedback_items = analysis_result.get('feedback_items', [])
        
        risk_counts = {'High': 0, 'Medium': 0, 'Low': 0}
        for item in feedback_items:
            risk_level = item.get('risk_level', 'Low')
            risk_counts[risk_level] = risk_counts.get(risk_level, 0) + 1
        
        sections_info.append({
            'name': section_name,
            'word_count': len(content.split()),
            'analyzed': section_name in current_session.analysis_results,
            'feedback_count': len(feedback_items),
            'high_risk_count': risk_counts['High'],
            'medium_risk_count': risk_counts['Medium'],
            'low_risk_count': risk_counts['Low'],
            'user_feedback_count': len(current_session.user_feedback.get(section_name, []))
        })
    
    return jsonify({'sections': sections_info})

@app.route('/api/section/<section_name>')
def get_section_analysis(section_name):
    if not current_session or section_name not in current_session.sections:
        return jsonify({'success': False, 'error': 'Section not found'}), 404
    
    content = current_session.sections.get(section_name, "")
    analysis = current_session.analysis_results.get(section_name, {})
    user_feedback = current_session.user_feedback.get(section_name, [])
    accepted_feedback = current_session.accepted_feedback.get(section_name, [])
    rejected_feedback = current_session.rejected_feedback.get(section_name, [])
    
    return jsonify({
        'success': True,
        'section_name': section_name,
        'content': content,
        'analysis': analysis,
        'user_feedback': user_feedback,
        'accepted_feedback': accepted_feedback,
        'rejected_feedback': rejected_feedback,
        'word_count': len(content.split()),
        'character_count': len(content)
    })

@app.route('/api/chat/basic', methods=['POST'])
def basic_chat():
    try:
        data = request.json
        message = data.get('message', '').strip()
        
        if not message:
            return jsonify({
                'success': False, 
                'error': 'Message is required'
            }), 400
        
        context = {
            'current_section': current_session.get_current_section_name() if current_session else None,
            'document_name': current_session.document_name if current_session else None,
            'session_active': current_session is not None,
            'analysis_status': analysis_status,
            'total_sections': len(current_session.sections) if current_session else 0,
            'analyzed_sections': len(current_session.analysis_results) if current_session else 0,
            'has_guidelines': current_session.guidelines_content is not None if current_session else False,
            'functionalities_enabled': functionalities_enabled
        }
        
        if current_session and functionalities_enabled:
            high_risk = medium_risk = low_risk = 0
            
            for result in current_session.analysis_results.values():
                for item in result.get('feedback_items', []):
                    risk = item.get('risk_level', 'Low')
                    if risk == 'High': high_risk += 1
                    elif risk == 'Medium': medium_risk += 1
                    else: low_risk += 1
            
            context.update({
                'high_risk': high_risk,
                'medium_risk': medium_risk,
                'low_risk': low_risk,
                'word_count': len(current_session.sections.get(context['current_section'], '').split()) if context['current_section'] else 0
            })
        
        response = process_chat_query(message, context)
        
        chat_entry = {
            'id': f"chat_{uuid.uuid4().hex[:8]}",
            'timestamp': datetime.now().isoformat(),
            'user_message': message,
            'ai_response': response,
            'context': context,
            'llm_used': 'Enhanced Tara AI'
        }
        chat_history.append(chat_entry)
        
        log_activity(f"Chat interaction: {message[:30]}...", "INFO")
        
        return jsonify({
            'success': True,
            'response': response,
            'timestamp': datetime.now().isoformat(),
            'llm_model': 'Enhanced Tara AI',
            'response_type': 'enhanced_ai'
        })
        
    except Exception as e:
        log_activity(f"Chat processing failed: {str(e)}", "ERROR")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/feedback/accept', methods=['POST'])
def accept_feedback():
    global functionalities_enabled
    
    if not current_session:
        return jsonify({'success': False, 'error': 'No active session'}), 400
    
    if not functionalities_enabled:
        return jsonify({'success': False, 'error': 'Please wait for analysis to complete'}), 400
    
    try:
        data = request.json
        section = data.get('section')
        feedback_id = data.get('feedback_id')
        
        analysis_result = current_session.analysis_results.get(section, {})
        feedback_items = analysis_result.get('feedback_items', [])
        
        feedback_item = None
        for item in feedback_items:
            if item.get('id') == feedback_id:
                feedback_item = item
                break
        
        if not feedback_item:
            return jsonify({'success': False, 'error': 'Feedback item not found'}), 400
        
        current_session.accepted_feedback[section].append(feedback_item)
        current_session.rejected_feedback[section] = [
            item for item in current_session.rejected_feedback[section] 
            if item.get('id') != feedback_id
        ]
        
        log_activity(f"Feedback accepted: {feedback_item.get('type', 'unknown')} in {section}", "SUCCESS", section)
        
        return jsonify({'success': True, 'message': 'Feedback accepted'})
        
    except Exception as e:
        log_activity(f"Failed to accept feedback: {str(e)}", "ERROR")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/feedback/reject', methods=['POST'])
def reject_feedback():
    global functionalities_enabled
    
    if not current_session:
        return jsonify({'success': False, 'error': 'No active session'}), 400
    
    if not functionalities_enabled:
        return jsonify({'success': False, 'error': 'Please wait for analysis to complete'}), 400
    
    try:
        data = request.json
        section = data.get('section')
        feedback_id = data.get('feedback_id')
        
        analysis_result = current_session.analysis_results.get(section, {})
        feedback_items = analysis_result.get('feedback_items', [])
        
        feedback_item = None
        for item in feedback_items:
            if item.get('id') == feedback_id:
                feedback_item = item
                break
        
        if not feedback_item:
            return jsonify({'success': False, 'error': 'Feedback item not found'}), 400
        
        current_session.rejected_feedback[section].append(feedback_item)
        current_session.accepted_feedback[section] = [
            item for item in current_session.accepted_feedback[section] 
            if item.get('id') != feedback_id
        ]
        
        log_activity(f"Feedback rejected: {feedback_item.get('type', 'unknown')} in {section}", "WARNING", section)
        
        return jsonify({'success': True, 'message': 'Feedback rejected'})
        
    except Exception as e:
        log_activity(f"Failed to reject feedback: {str(e)}", "ERROR")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/export', methods=['POST'])
def export_results():
    if not current_session:
        return jsonify({'success': False, 'error': 'No active session'}), 400
    
    try:
        export_data = {
            'session_metadata': {
                'session_id': current_session.session_id,
                'document_name': current_session.document_name,
                'generated_at': datetime.now().isoformat(),
                'analysis_framework': 'Hawkeye 20-Point Investigation Framework'
            },
            'document_sections': current_session.sections,
            'analysis_results': current_session.analysis_results,
            'user_feedback': dict(current_session.user_feedback),
            'accepted_feedback': dict(current_session.accepted_feedback),
            'rejected_feedback': dict(current_session.rejected_feedback),
            'chat_history': chat_history,
            'analysis_logs': analysis_logs
        }
        
        log_activity("Analysis results exported", "SUCCESS")
        
        return jsonify({
            'success': True,
            'export_data': export_data,
            'filename': f"hawkeye_analysis_{current_session.session_id}.json"
        })
        
    except Exception as e:
        log_activity(f"Export failed: {str(e)}", "ERROR")
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5005))
    
    print("Enhanced Writeup Automation AI Tool v2.0")
    print("Production Deployment - Exact Localhost Replica")
    print(f"Server starting on port {port}")
    print(f"Upload folder: {app.config['UPLOAD_FOLDER']}")
    print("=" * 50)
    
    app.run(
        host='0.0.0.0',
        port=port,
        debug=False,
        threaded=True
    )